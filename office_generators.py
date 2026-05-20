"""Generadores de Word (.docx) y Excel (.xlsx) con branding del cliente.

Ambas funciones reciben datos estructurados (sections/sheets) y devuelven
{filename, url, size_bytes}. Diseñados para ser tolerantes a entradas malformadas
del LLM (acepta JSON-string o dict/list nativos).
"""

from __future__ import annotations

import json
import secrets
from datetime import datetime
from pathlib import Path
from typing import Any

BRAND_HEX  = "1A4EA0"  # personalizar por instancia si querés otro color
BRAND_RGB  = (0x1A, 0x4E, 0xA0)
TEXT_HEX   = "1D1D1F"
SUBT_HEX   = "515154"


def _company_name() -> str:
    """Lee el nombre desde settings (no hardcoded)."""
    try:
        from app.config import settings
        return settings.company_full_name
    except Exception:
        return "Demo Company S.A."


def _safe_filename(prefix: str, title: str, ext: str) -> str:
    base = "".join(c if c.isalnum() else "_" for c in (title or "documento"))
    base = (base.strip("_") or "documento")[:40]
    rand = secrets.token_hex(3)
    return f"{prefix}_{base}_{rand}.{ext}"


def _parse_maybe_json(value: Any, default):
    """Acepta el valor como JSON-string, o ya parseado. Devuelve default si falla."""
    if value is None:
        return default
    if isinstance(value, (list, dict)):
        return value
    if isinstance(value, str):
        s = value.strip()
        if not s:
            return default
        try:
            return json.loads(s)
        except Exception:
            return default
    return default


def _normalize_sections(value: Any) -> list[dict]:
    """Acepta secciones en múltiples formatos y devuelve list[dict] estándar."""
    parsed = _parse_maybe_json(value, None)
    if parsed is None:
        return []
    if isinstance(parsed, list):
        return [s for s in parsed if isinstance(s, dict)] + \
               [{"type": "paragraph", "text": str(s)} for s in parsed if not isinstance(s, dict) and s]
    if isinstance(parsed, dict):
        if "sections" in parsed and isinstance(parsed["sections"], list):
            return _normalize_sections(parsed["sections"])
        # Probablemente una sola sección
        if "type" in parsed:
            return [parsed]
        return []
    if isinstance(parsed, str) and parsed.strip():
        return [{"type": "paragraph", "text": p.strip()} for p in parsed.split("\n\n") if p.strip()]
    return []


def _records_to_sheet_dict(records: list, name: str) -> dict:
    """Convierte una lista de dicts en {name, headers, rows} preservando orden de columnas."""
    seen: list = []
    for r in records:
        if isinstance(r, dict):
            for k in r.keys():
                if k not in seen:
                    seen.append(k)
    rows = [[r.get(c, "") if isinstance(r, dict) else "" for c in seen] for r in records]
    return {"name": name, "headers": seen, "rows": rows}


def _normalize_one_sheet(sheet: Any, idx: int) -> dict:
    """Normaliza una sola hoja a {name, headers, rows} con rows como list[list]."""
    if not isinstance(sheet, dict):
        return {"name": f"Hoja{idx}", "headers": [], "rows": []}

    name    = str(sheet.get("name") or sheet.get("title") or sheet.get("sheet") or f"Hoja{idx}")
    headers = sheet.get("headers") or sheet.get("header") or sheet.get("columns") or []
    rows    = sheet.get("rows") or sheet.get("data") or sheet.get("values") or []

    if isinstance(headers, str):
        headers = [headers]
    if not isinstance(headers, list):
        headers = []

    # Rows en formato column-oriented: {"col1": [v1, v2], "col2": [v3, v4]}
    if isinstance(rows, dict):
        cols = list(rows.keys())
        max_len = max((len(v) for v in rows.values() if isinstance(v, list)), default=0)
        if not headers:
            headers = cols
        flat = []
        for i in range(max_len):
            flat.append([(rows[c][i] if isinstance(rows[c], list) and i < len(rows[c]) else "") for c in cols])
        rows = flat
    # Rows como lista
    elif isinstance(rows, list):
        # Caso: lista de dicts (records) → convertir a list de lists
        if rows and isinstance(rows[0], dict):
            if not headers:
                seen = []
                for r in rows:
                    if isinstance(r, dict):
                        for k in r.keys():
                            if k not in seen:
                                seen.append(k)
                headers = seen
            rows = [[r.get(c, "") if isinstance(r, dict) else "" for c in headers] for r in rows]
        else:
            # Normalizar cada fila a lista
            rows = [r if isinstance(r, list) else ([r] if r is not None else []) for r in rows]
    else:
        rows = []

    return {"name": name, "headers": list(headers), "rows": rows}


def _normalize_sheets(value: Any) -> list[dict]:
    """Acepta sheets en múltiples formatos y devuelve list[{name, headers, rows}]."""
    parsed = _parse_maybe_json(value, None)
    if parsed is None:
        return []

    # Caso: dict envuelto, {"sheets": [...]}
    if isinstance(parsed, dict) and isinstance(parsed.get("sheets"), list):
        return _normalize_sheets(parsed["sheets"])

    # Caso: dict que ES una sola hoja (tiene name/rows/headers/data)
    if isinstance(parsed, dict) and any(k in parsed for k in ("rows", "data", "headers", "values")):
        return [_normalize_one_sheet(parsed, 1)]

    # Caso: dict de hojas con nombre como key: {"Hoja1": {...}, "Hoja2": {...}}
    if isinstance(parsed, dict):
        out = []
        for i, (k, v) in enumerate(parsed.items(), 1):
            if isinstance(v, dict):
                merged = {**v, "name": v.get("name") or str(k)}
                out.append(_normalize_one_sheet(merged, i))
            elif isinstance(v, list):
                out.append(_normalize_one_sheet({"name": str(k), "rows": v}, i))
        return out

    if isinstance(parsed, list):
        if not parsed:
            return []
        first = parsed[0]
        # Caso: list de sheets ({name, rows, ...})
        if isinstance(first, dict) and any(k in first for k in ("rows", "data", "headers", "name", "values")):
            return [_normalize_one_sheet(s, i) for i, s in enumerate(parsed, 1)]
        # Caso: list de records (sin wrapper de sheet) — una sola hoja
        if isinstance(first, dict):
            return [_records_to_sheet_dict(parsed, "Hoja1")]
        # Caso: list de listas — una sola hoja sin headers
        if isinstance(first, list):
            return [{"name": "Hoja1", "headers": [], "rows": parsed}]

    return []


# ──────────────────────────────────────────────────────────────
# WORD
# ──────────────────────────────────────────────────────────────

def generate_word_document(
    out_dir: Path,
    *,
    title: str,
    sections: Any,
    subtitle: str = "",
    author: str | None = None,
) -> dict:
    if author is None:
        author = _company_name()
    """Genera un .docx con secciones tipadas.

    `sections` puede venir como list o como JSON-string. Cada item tiene `type`:
      - {"type": "heading", "level": 1|2|3, "text": "..."}
      - {"type": "paragraph", "text": "..."}
      - {"type": "bullet", "items": ["...", "..."]}
      - {"type": "numbered", "items": ["...", "..."]}
      - {"type": "table", "rows": [["c1","c2"], ["v1","v2"]], "header": true}
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sections_list = _normalize_sections(sections)
    if not sections_list:
        raise ValueError(
            "No se recibieron secciones válidas para el Word. "
            "Pasá sections como JSON-string con una lista de objetos {type, ...}. "
            "Ejemplo: [{\"type\":\"heading\",\"level\":1,\"text\":\"Título\"},{\"type\":\"paragraph\",\"text\":\"Texto\"}]"
        )

    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Banner / título
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = head.add_run(_company_name())
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*BRAND_RGB)
    run.bold = True

    t = doc.add_paragraph()
    rt = t.add_run(title or "Documento")
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(*BRAND_RGB)

    if subtitle:
        s = doc.add_paragraph()
        rs = s.add_run(subtitle)
        rs.font.size = Pt(12)
        rs.font.color.rgb = RGBColor(0x51, 0x51, 0x54)
        rs.italic = True

    # Fecha
    meta = doc.add_paragraph()
    rm = meta.add_run(datetime.now().strftime("%d de %B de %Y"))
    rm.font.size = Pt(9)
    rm.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)

    # Separador
    doc.add_paragraph("―" * 30)

    # Secciones
    for item in sections_list:
        if not isinstance(item, dict):
            doc.add_paragraph(str(item))
            continue
        kind = (item.get("type") or "").lower().strip()
        if kind == "heading":
            level = max(1, min(3, int(item.get("level", 1) or 1)))
            p = doc.add_paragraph()
            r = p.add_run(item.get("text", ""))
            r.bold = True
            r.font.size = Pt({1: 16, 2: 14, 3: 12}[level])
            r.font.color.rgb = RGBColor(*BRAND_RGB)
        elif kind == "paragraph":
            doc.add_paragraph(item.get("text", ""))
        elif kind == "bullet":
            for li in (item.get("items") or []):
                doc.add_paragraph(str(li), style="List Bullet")
        elif kind == "numbered":
            for li in (item.get("items") or []):
                doc.add_paragraph(str(li), style="List Number")
        elif kind == "table":
            norm = _normalize_one_sheet(
                {"rows": item.get("rows"), "headers": item.get("headers")},
                1,
            )
            data_rows    = norm["rows"]
            norm_headers = norm["headers"]
            has_header   = bool(item.get("header", True))
            final_rows: list = []
            if norm_headers:
                final_rows.append(norm_headers)
                has_header = True
            final_rows.extend(data_rows)
            if not final_rows:
                continue
            ncols = max(len(r) for r in final_rows)
            tbl = doc.add_table(rows=len(final_rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for ri, row in enumerate(final_rows):
                for ci in range(ncols):
                    val = row[ci] if ci < len(row) else ""
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = str(val)
                    if has_header and ri == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        else:
            # Tipo desconocido: ignorar o renderizar como párrafo
            text = item.get("text") or json.dumps(item, ensure_ascii=False)
            doc.add_paragraph(text)

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    rf = foot.add_run(f"Generado por {author} — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0x86, 0x86, 0x8B)
    rf.italic = True

    filename = _safe_filename("cortex", title, "docx")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    doc.save(out_path)
    return {
        "filename":   filename,
        "url":        f"/documents/{filename}",
        "size_bytes": out_path.stat().st_size,
    }


# ──────────────────────────────────────────────────────────────
# EXCEL
# ──────────────────────────────────────────────────────────────

def generate_excel_spreadsheet(
    out_dir: Path,
    *,
    title: str,
    sheets: Any,
) -> dict:
    """Genera un .xlsx con una o varias hojas.

    `sheets` puede venir como list o como JSON-string. Cada hoja:
      {"name": "Hoja1", "headers": ["col1","col2"], "rows": [["v1","v2"], ...]}
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    sheets_list = _normalize_sheets(sheets)
    if not sheets_list:
        raise ValueError(
            "No se recibieron hojas válidas para el Excel. "
            "Pasá sheets como JSON-string. Formatos válidos: "
            "(1) lista de hojas: [{\"name\":\"H1\",\"headers\":[\"c1\",\"c2\"],\"rows\":[[\"v1\",\"v2\"]]}]. "
            "(2) hoja única: {\"name\":\"H1\",\"headers\":[...],\"rows\":[[...]]}. "
            "(3) lista de records: [{\"col1\":\"v1\",\"col2\":\"v2\"}, ...]."
        )
    # Validar que al menos UNA hoja tenga datos
    total_cells = sum(
        (len(s.get("headers") or []) + sum(len(r) for r in (s.get("rows") or [])))
        for s in sheets_list
    )
    if total_cells == 0:
        raise ValueError("Las hojas están vacías (sin headers ni rows). Pasá los datos en `rows` o `headers`.")

    wb = Workbook()
    # Borrar la hoja por defecto si vamos a crear las nuestras
    wb.remove(wb.active)

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor=BRAND_HEX)
    align       = Alignment(vertical="center", wrap_text=False)
    thin        = Side(border_style="thin", color="DDDDDD")
    border      = Border(left=thin, right=thin, top=thin, bottom=thin)

    for idx, sheet in enumerate(sheets_list, 1):
        if not isinstance(sheet, dict):
            continue
        name = str(sheet.get("name") or f"Hoja{idx}")[:31]  # Excel limit
        # Sanear nombre: caracteres prohibidos
        for c in r'[]:*?/\\':
            name = name.replace(c, "_")
        ws = wb.create_sheet(title=name)

        headers = sheet.get("headers") or []
        rows    = sheet.get("rows") or []

        row_offset = 0
        if headers:
            for ci, h in enumerate(headers, 1):
                cell = ws.cell(row=1, column=ci, value=str(h))
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = align
                cell.border = border
            row_offset = 1

        for ri, row in enumerate(rows, 1 + row_offset):
            if not isinstance(row, list):
                continue
            for ci, val in enumerate(row, 1):
                cell = ws.cell(row=ri, column=ci, value=val if val is None else (val if isinstance(val, (int, float, bool)) else str(val)))
                cell.border = border
                cell.alignment = align

        # Auto-ajuste rudimentario de ancho de columna
        max_col = max(len(headers), max((len(r) for r in rows if isinstance(r, list)), default=0))
        for ci in range(1, max_col + 1):
            letter = get_column_letter(ci)
            max_len = 8
            if headers and ci <= len(headers):
                max_len = max(max_len, len(str(headers[ci-1])))
            for r in rows[:200]:  # mirar las primeras 200 filas para inferir ancho
                if isinstance(r, list) and ci <= len(r):
                    max_len = max(max_len, min(40, len(str(r[ci-1]))))
            ws.column_dimensions[letter].width = max_len + 2

    filename = _safe_filename("cortex", title, "xlsx")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    wb.save(out_path)
    return {
        "filename":   filename,
        "url":        f"/documents/{filename}",
        "size_bytes": out_path.stat().st_size,
    }
